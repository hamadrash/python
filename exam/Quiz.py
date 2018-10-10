from docx import Document
from random import choice

numbers = [x for x in range(2, 120) if x % 2 == 0 ]

sum1 = choice(numbers)
sum2 = choice(numbers)


sub1 = choice(numbers)
sub2 = choice([x for x in range(2,120) if x < sub1])

multi1 = choice(numbers)
multi2 = choice(numbers)

div1 = choice(numbers)
div2 = choice([x for x in range(2 ,120) if div1 % x == 0]) 
dot = "."*70
d = Document()
d.add_heading('AlRowad Private School - Math Exam 1', 0)
d.add_heading("Name : .............................. Grade : .............", 1 )

d.add_heading('Answer The Folowing Question :\n\n', 1)
d.add_paragraph(str(sum1)+"+"+str(sum2)+dot+"\n\n" , style='ListNumber')
d.add_paragraph(str(sub1)+"-"+str(sub2)+dot+"\n\n" , style='ListNumber')
d.add_paragraph(str(multi1)+"x"+str(multi2)+dot+"\n\n" , style='ListNumber')
d.add_paragraph(str(div1)+"/"+str(div2)+dot+"\n\n" , style='ListNumber')

d.add_heading("End of The Exam.", 1)

d.add_page_break()
d.save('test1.docx')